import io
import re

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import db, models

router = APIRouter(prefix="/api/drafts", tags=["drafts"])

FORUM_LABELS = {
    "lokayuktha": "Lokayuktha",
    "lokayukta": "Lokayuktha",
    "nhrc": "NHRC",
    "womens_commission": "State Women's Commission",
    "rti": "RTI Application",
    "consumer_forum": "Consumer Forum",
    "district_court": "District Court",
}


def _forum_label(value):
    if not value:
        return "Unspecified Forum"
    return FORUM_LABELS.get(value, value.replace("_", " ").title())


# --- Defensive markdown stripping -------------------------------------------
# The model is instructed to never emit markdown, but instructions get
# ignored. Nothing below may ever reach the .docx unprocessed.

_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_HEADING_MARKER_PATTERN = re.compile(r"^\s{0,3}#{1,6}\s*", re.MULTILINE)
_DIVIDER_LINE_PATTERN = re.compile(r"^\s*([-*_])\1{2,}\s*$", re.MULTILINE)
_TABLE_ROW_PATTERN = re.compile(r"^\s*\|.*\|\s*$", re.MULTILINE)


def _clean_table_row(match: re.Match) -> str:
    replaced = match.group(0).replace("|", " ").strip()
    # Pure separator rows like "|---|---|" collapse to dashes/colons only —
    # drop them entirely rather than leaving stray dashes behind.
    if re.fullmatch(r"[-:\s]*", replaced):
        return ""
    return replaced


def _clean_draft_text(text: str) -> str:
    text = _HEADING_MARKER_PATTERN.sub("", text)
    text = _DIVIDER_LINE_PATTERN.sub("", text)
    text = _TABLE_ROW_PATTERN.sub(_clean_table_row, text)
    return text


def _split_bold_runs(line: str):
    """Split a line into (text, is_bold) segments, consuming ** markers.

    Bold segments render as bold runs; the ** markers themselves never
    appear in the output either way.
    """
    segments = []
    pos = 0
    for m in _BOLD_PATTERN.finditer(line):
        if m.start() > pos:
            segments.append((line[pos:m.start()], False))
        segments.append((m.group(1), True))
        pos = m.end()
    if pos < len(line):
        segments.append((line[pos:], False))
    if not segments:
        segments.append((line, False))
    return segments


def _is_heading(line: str) -> bool:
    stripped = line.strip().strip(":")
    if not stripped or len(stripped) > 70:
        return False
    return stripped.isupper()


def _has_signature_block(content: str) -> bool:
    lower = content.lower()
    return "place:" in lower and "date:" in lower and "signature" in lower


@router.get("/{draft_id}/download")
def download_draft(draft_id: int, db_session: Session = Depends(db.get_db)):
    draft = db_session.query(models.Draft).filter(models.Draft.id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    case = db_session.query(models.Case).filter(models.Case.id == draft.case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    doc = DocxDocument()
    cleaned = _clean_draft_text(draft.content)

    for raw_line in cleaned.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        if _is_heading(line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace("*", ""))
            run.bold = True
            continue

        # Body text: justified, model's own numbering/text kept literal —
        # never wrapped in a Word auto-numbered list style.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for segment_text, is_bold in _split_bold_runs(line):
            if not segment_text:
                continue
            run = p.add_run(segment_text)
            run.bold = is_bold

    # Safety net only — skip if the draft already produced its own block,
    # since the model is now instructed to include one itself.
    if not _has_signature_block(draft.content):
        doc.add_paragraph()
        for line in [
            "Place: ______________________",
            "Date: ______________________",
            "",
            "Signature: ______________________",
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", case.title)[:60] or "draft"
    filename = f"{safe_title}_draft.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
